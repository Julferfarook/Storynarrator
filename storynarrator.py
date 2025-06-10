import streamlit as st
import google.generativeai as genai
from docx import Document
from fpdf import FPDF
from unidecode import unidecode  # ✅ Add this
import os

# --- CONFIGURE GEMINI FLASH ---
genai.configure(api_key="AIzaSyCsINRQVql0DWC3uZiWmA47ZLwS-EGkdPE")  # Replace with your actual API key
model = genai.GenerativeModel("models/gemini-2.0-flash")

# --- STREAMLIT PAGE SETUP ---
st.set_page_config(page_title="📚 Story Narrator", page_icon="📘")
st.title("📘 AI Story Narrator")
st.markdown("Generate imaginative stories using Gemini 2.0 Flash and download them in PDF or Word format.")

# --- INPUT FIELDS ---
topic = st.text_input("📝 Enter story topic", placeholder="e.g. The Lost Kingdom of Atlantis")
word_count = st.text_input("🔢 Desired word count", placeholder="e.g. 300")

story = ""

# --- GENERATE STORY ---
if st.button("🚀 Generate Story"):
    if not topic or not word_count.isdigit():
        st.warning("Please enter a valid topic and a numeric word count.")
    else:
        with st.spinner("Generating story using Gemini Flash..."):
            prompt = f"Write a creative story of about {word_count} words on the topic: {topic}."
            response = model.generate_content(prompt)
            story = response.text.strip()
            st.success("Story generated!")
            st.text_area("📖 Your Story", value=story, height=300)

            # ✅ Convert story text to ASCII to avoid UnicodeEncodeError
            story_ascii = unidecode(story)

            # --- GENERATE PDF ---
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Arial", size=12)
            for line in story_ascii.split('\n'):
                pdf.multi_cell(0, 10, line)
            pdf_path = "story.pdf"
            pdf.output(pdf_path)

            # --- GENERATE WORD DOCX ---
            doc = Document()
            doc.add_heading("Generated Story", 0)
            doc.add_paragraph(story)  # Keep original Unicode in Word
            word_path = "story.docx"
            doc.save(word_path)

            # --- DOWNLOAD BUTTONS ---
            col1, col2 = st.columns(2)
            with col1:
                with open(pdf_path, "rb") as f:
                    st.download_button(
                        label="📄 Download as PDF",
                        data=f,
                        file_name="story.pdf",
                        mime="application/pdf"
                    )
            with col2:
                with open(word_path, "rb") as f:
                    st.download_button(
                        label="📝 Download as Word",
                        data=f,
                        file_name="story.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            # Optional cleanup
            os.remove(pdf_path)
            os.remove(word_path)
